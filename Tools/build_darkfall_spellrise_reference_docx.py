from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path("Docs/References/DarkfallManual")
SRC_JSON = OUT_DIR / "darkfall_manual_pages.json"
DOCX_PATH = OUT_DIR / "Darkfall_Unholy_Wars_Referencia_para_SpellRise.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 90, 90)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALIBRI = "Calibri"


def set_run_font(run, name=CALIBRI, size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_grid(table, widths):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = tbl.find(qn("w:tblGrid"))
    if grid is not None:
        tbl.remove(grid)
    grid = OxmlElement("w:tblGrid")
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    tbl.insert(1, grid)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def paragraph(doc, text="", style=None, before=None, after=None, bold_label=None):
    p = doc.add_paragraph(style=style)
    if before is not None:
        p.paragraph_format.space_before = Pt(before)
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    if bold_label and text.startswith(bold_label):
        r1 = p.add_run(bold_label)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_label):])
        set_run_font(r2)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def bullet(doc, text):
    p = paragraph(doc, text, style="List Bullet", after=4)
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    return p


def numbered(doc, text):
    p = paragraph(doc, text, style="List Number", after=4)
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    return p


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.line_spacing = 1.25
    if level == 1:
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(10)
    elif level == 2:
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(7)
    else:
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(5)
    for run in p.runs:
        set_run_font(run, size={1: 16, 2: 13, 3: 12}.get(level, 11),
                     color=BLUE if level < 3 else DARK_BLUE, bold=True)
    return p


def add_callout(doc, title, items):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_grid(table, [9360])
    cell = table.cell(0, 0)
    shade_cell(cell, LIGHT_GRAY)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_run_font(r, bold=True, color=DARK_BLUE)
    for item in items:
        p2 = cell.add_paragraph(style="List Bullet")
        p2.paragraph_format.space_after = Pt(2)
        p2.paragraph_format.left_indent = Inches(0.25)
        p2.paragraph_format.first_line_indent = Inches(-0.125)
        set_run_font(p2.add_run(item), size=10.5)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_grid(table, widths)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        shade_cell(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(header)
        set_run_font(r, size=10.5, bold=True, color=DARK_BLUE)
    for row_data in rows:
        row = table.add_row()
        for idx, value in enumerate(row_data):
            cell = row.cells[idx]
            set_cell_width(cell, widths[idx])
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(value)
            set_run_font(r, size=10)
    set_table_grid(table, widths)
    paragraph(doc, "", after=4)
    return table


def setup_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    normal = doc.styles["Normal"]
    normal.font.name = CALIBRI
    normal._element.rPr.rFonts.set(qn("w:ascii"), CALIBRI)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), CALIBRI)
    normal.font.size = Pt(11)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = CALIBRI
        style._element.rPr.rFonts.set(qn("w:ascii"), CALIBRI)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), CALIBRI)
    header = section.header.paragraphs[0]
    header.text = "SpellRise | Referencia externa de design"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run_font(header.runs[0], size=9, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.text = "Uso interno - sintese baseada no manual publico de Darkfall Unholy Wars"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(footer.runs[0], size=9, color=MUTED)


def build_doc():
    pages = json.loads(SRC_JSON.read_text(encoding="utf-8"))
    source_map = {p["title"]: p["url"] for p in pages}
    doc = Document()
    setup_styles(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(10)
    title.paragraph_format.space_after = Pt(4)
    r = title.add_run("Darkfall Unholy Wars")
    set_run_font(r, size=24, color=RGBColor(0, 0, 0), bold=True)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(16)
    set_run_font(subtitle.add_run("Resumo do manual e referencia aplicada para SpellRise"), size=14, color=MUTED)

    meta = [
        ("Finalidade", "Guia de referencia para design e producao de sistemas hardcore PvP, full loot e free target."),
        ("Fonte", "Manual publico de Darkfall Unholy Wars, coletado de http://www.darkfallonline.com/manual."),
        ("Data", "27/05/2026."),
        ("Uso", "Inspiracao e comparacao tecnica. Nao e especificacao direta nem copia de regras."),
    ]
    add_table(doc, ["Campo", "Valor"], meta, [1900, 7460])

    add_callout(doc, "Leitura executiva para SpellRise", [
        "A referencia mais forte e a combinacao entre combate manual, consequencia real de morte, economia baseada em itens e organizacao social via clans.",
        "A progressao funciona melhor quando atividades diferentes convergem para uma moeda/XP de longo prazo, mas crafting preserva identidade por uso e material consumido.",
        "A maior adaptacao tecnica para SpellRise e transformar cada acao sensivel em fluxo server-authoritative: input local, validacao no servidor, commit, replicacao minima e apresentacao local.",
    ])

    heading(doc, "1. Principios aproveitaveis", 1)
    for item in [
        "Combate deve depender de mira, leitura de movimento, stamina e decisao de timing, nao de selecao automatica de alvo.",
        "Morte precisa gerar perda concreta e oportunidade de recuperacao, criando risco de rota, inventario e decisao de engajamento.",
        "Progressao horizontal e troca de papel mantem variedade sem obrigar reroll, mas precisa de custo e janela de vulnerabilidade para nao virar troca instantanea sem compromisso.",
        "Clans, holdings, cercos e vilas transformam PvP avulso em disputa territorial com calendario, custo, premio e risco politico.",
        "A economia ganha densidade quando crafting, vendors, bancos, casas, loot e transporte conversam entre si e quando itens carregados podem ser perdidos.",
    ]:
        bullet(doc, item)

    heading(doc, "2. Matriz de adaptacao para SpellRise", 1)
    add_table(doc, ["Sistema Darkfall", "Ideia central", "Uso recomendado em SpellRise"], [
        ("Prowess", "Moeda de progressao obtida por atividades variadas.", "Usar como referencia para XP ampla, feats e objetivos. Gastos autoritativos no servidor e auditaveis."),
        ("Roles e escolas", "Papel flexivel com duas especializacoes e restricoes de equipamento.", "Aplicar como loadouts/classes data-driven, mantendo ASC no PlayerState e sem depender de VIG/FOC."),
        ("Combate manual", "Ataques e abilities por mira, cast, canal e cooldown.", "Mapear para GAS: ativacao local, target data validada, commit server, GE autoritativo e cues replicadas."),
        ("Tombstone/full loot", "Morte cria recipiente lootavel e retorno ao local.", "Projetar morte/loot server-side, com duracao, ownership temporario e reconciliacao anti-dupe."),
        ("Safe zones", "Areas iniciais impedem dano e roubo; fora delas o risco e pleno.", "Usar volumes server-authoritative com tags de area e bloqueio de dano/loot no servidor."),
        ("Clans/holdings", "Bases de clan, estruturas, permissões, bindstone e conflito territorial.", "Integrar com building mode e persistencia, com RPCs limitados e matriz de permissao server-side."),
        ("Sieges", "Desafio com fases, custo, janela de preparacao e objetivos vulneraveis.", "Basear guerra territorial em estado persistido, cronograma verificavel e objetivos replicados com baixo payload."),
        ("Villages", "Captura e roubo periodico de recursos em area contestavel.", "Excelente referencia para hotspots PvP de baixo custo operacional e recompensa previsivel."),
        ("Crafting", "Maioria dos itens vem de jogadores; falha consome materiais.", "Aplicar economia player-driven com recipes no servidor, auditoria e protecao contra rollback/dupe."),
        ("Teleportes", "Mobilidade de longa distancia custa recurso e cria risco de chegada.", "Evitar mobilidade gratuita; todo teleporte deve ter custo, cast/canal e validacao server-side."),
    ], [1700, 3100, 4560])

    heading(doc, "3. Progressao, atributos e habilidades", 1)
    paragraph(doc, "Darkfall usa Prowess como acumulador de experiencia de mundo: combate PvE, crafting, harvesting, feitos, vilas e acoes sociais/economicas alimentam uma moeda de progressao. O gasto desbloqueia atributos, habilidades e boosters. Crafting e excecao importante: evolui pelo proprio uso e pelos materiais consumidos.")
    paragraph(doc, "Para SpellRise, a ideia mais util e separar progresso de personagem, especializacao e oficio. STR, AGI, INT e WIS ja encaixam melhor que os aliases legados. Qualquer custo, compra de skill, booster ou respec deve ser commitado no servidor, com log de auditoria e rollback seguro.")
    add_callout(doc, "Aplicacao pratica", [
        "Feats podem guiar jogador novo sem virar questline obrigatoria.",
        "Boosters devem ser data assets com clamps de balance e sem dependencia de VIG/FOC.",
        "Crafting deve ter chance, qualidade, autoria do maker e consumo server-side de materiais.",
    ])

    heading(doc, "4. Combate e UX de habilidade", 1)
    paragraph(doc, "O manual reforca combate em tempo real, rapido e manualmente mirado. Abilities podem ser instantaneas, ter cast, serem pre-cast por botao segurado, aplicar efeito ao longo do tempo ou canalizar enquanto o jogador mantem input. Health, stamina e mana sao recursos centrais; stamina tambem controla sprint e salto.")
    paragraph(doc, "Em SpellRise, isso deve virar contrato GAS claro: input local apenas expressa intencao; o servidor valida target data, alcance, linha de visao, custo, cooldown e ownership; o commit acontece no servidor; GameplayEffects resolvem dano/recurso; GameplayCues e atores replicados cuidam da apresentacao.")
    add_table(doc, ["Fase", "Contrato recomendado"], [
        ("Inicio", "Cliente coleta input, mira local e escolhe ability/slot. Sem decisao final de hit."),
        ("Target data", "Payload pequeno, com limites por campo, timestamp e validacao de contexto."),
        ("Commit", "Servidor confirma custo/cooldown/tags antes de spawnar projetil ou aplicar GE."),
        ("Resultado", "Dano, morte, queda, revive, gank e loot sempre server-authoritative."),
        ("Apresentacao", "Cliente pode prever animacao/VFX, mas reconcilia com confirmacao e OnRep."),
    ], [1800, 7560])

    heading(doc, "5. Morte, gank, revive e full loot", 1)
    paragraph(doc, "Darkfall cria uma etapa intermediaria: o jogador derrotado fica vulneravel, pode ser finalizado por gank ou levantado por revive. Tombstones concentram os itens do morto e podem ser saqueados; alguns itens iniciais sao protegidos. Isso gera drama sem remover imediatamente toda possibilidade de interacao.")
    paragraph(doc, "Para SpellRise, a recomendacao e tratar morte como maquina de estados no servidor: Alive, Downed, Executed, Lootable, Respawning. O cliente apenas mostra UI e prompts. Loot em tumulo precisa de lock transacional, tempo de vida, permissao inicial se houver, log de item e protecao anti-dupe.")

    heading(doc, "6. Mundo, safe zones e interacao", 1)
    paragraph(doc, "O mundo e dividido entre areas protegidas e areas abertas. Safe zones bloqueiam dano e roubo; fora delas o risco de PvP, perda de montaria e saque volta a existir. Interacoes usam uma tecla comum: banco, bindstone, portal, tumulo, container, montaria, barco, comida, potion e deployables.")
    paragraph(doc, "Em SpellRise, interacao generica deve ser conveniente no cliente, mas a resolucao deve ficar no servidor. Volumes de safe zone precisam bloquear dano, target data ofensivo, loot ilegal e spawn abusivo no lado autoritativo. Banco e inventario persistente devem ficar fora do PlayerController.")

    heading(doc, "7. Inventario, economia e crafting", 1)
    paragraph(doc, "O manual descreve backpack com peso, ouro carregado, stacks, bags, bank vault, vendors, recursos, recipes, crafting stations e chance de falha. A maioria dos itens relevantes vem de jogadores, e itens exibem o maker. Essa combinacao sustenta economia com identidade social.")
    add_table(doc, ["Elemento", "Valor para SpellRise", "Risco principal"], [
        ("Peso/capacidade", "Forca escolhas antes de sair da base.", "Replicar inventario completo ou peso por item sem budget."),
        ("Bags/sacks", "Organizacao e loot com valor operacional.", "Nested containers podem gerar dupe se a persistencia nao for transacional."),
        ("Maker label", "Cria reputacao de crafter e comercio emergente.", "Precisa persistir autoria sem permitir spoof."),
        ("Falha de crafting", "Sink de materiais e progressao de oficio.", "Deve ser resolvida e auditada no servidor."),
        ("Vendors", "Treinamento, compra e venda basica.", "Preco/estoque nao pode vir do cliente."),
    ], [1900, 3900, 3560])

    heading(doc, "8. Social, party, chat e clans", 1)
    paragraph(doc, "Darkfall usa interacao direta por radial, party com marcadores e divisao de rewards, friends, bloqueio, whisper, canais de chat, clans com permissoes, politica clan-player e relacoes entre clans. O ponto forte e que a estrutura social tambem vira regra de gameplay: bind, permissao de holding, guerra e aliados.")
    paragraph(doc, "Para SpellRise, party e clan devem ser servicos server-side com replicacao enxuta para UI. Permissoes de clan nao devem ser cache autoritativo no cliente. Chat pode ser separado do runtime de combate para evitar payload concorrendo com dados net-critical.")

    heading(doc, "9. Territorio, cercos, vilas e mercenarios", 1)
    paragraph(doc, "O sistema de challenge de Darkfall combina custo, aposta, Dominion, fases temporais, siege stones, vulnerabilidade de atacantes e defensores, condicoes de vitoria e grace period. Isso reduz ataques gratuitos e cria planejamento. Vilas oferecem receita periodica, roubo contestavel e janelas de recaptura.")
    paragraph(doc, "Para SpellRise, essa e a referencia mais forte para live service: objetivos territoriais precisam de cronograma, custo de entrada, premio, janela de vulnerabilidade, estado persistido e logs. Mercenarios e contratos adicionam politica e gold sink, mas exigem regras server-side para bond, pagamento e elegibilidade.")

    heading(doc, "10. Mobilidade, mapa e descoberta", 1)
    paragraph(doc, "Mapa com filtros, fog of war, waypoints, marcadores pessoais, tumulo marcado e pontos de interesse reduz friccao sem eliminar risco. Teleportation chambers e summon friend encurtam distancias, mas exigem shard, canal, aceite e vulnerabilidade durante o processo.")
    paragraph(doc, "Para SpellRise, mobilidade rapida deve ter custo material, cast/canal cancelavel e validacao de destino. Recall para bindstone/casa e util, mas deve respeitar combate, cooldown, safe zone, estado de morte e economia de shards.")

    heading(doc, "11. Checklist de producao para adaptar ao SpellRise", 1)
    checks = [
        "Toda mutacao de inventario, gold, Prowess equivalente, crafting, morte, loot e territorio acontece no servidor.",
        "Nenhum fluxo de combate depende de HUD, widget, camera ou input local no Dedicated Server.",
        "Abilities seguem: input local -> target data -> validacao server -> commit -> GE/projetil server -> cue/apresentacao.",
        "PlayerController nao armazena estado autoritativo; usa RPC validado apenas como ponte de intencao.",
        "RPCs de building, loot, trade, crafting, party, clan e siege declaram payload maximo, ownership, contexto, rate-limit e rejeicao com log.",
        "OnRep reconcilia UI/apresentacao e nunca decide dano, custo, cooldown, morte, loot ou ownership.",
        "Sistemas persistentes tem auditoria contra dupe, replay, rollback indevido e concorrencia de salvamento.",
        "Toda feature territorial declara budget de rede: atores replicados, propriedades, condicao, NetUpdateFrequency, RPC/s e risco de burst.",
        "Mudancas sensiveis entram com default seguro no servidor e caminho de rollback via config/data asset quando viavel.",
    ]
    for check in checks:
        numbered(doc, check)

    heading(doc, "12. O que nao copiar diretamente", 1)
    for item in [
        "Nao copiar UI radial ou controles de Darkfall literalmente; adaptar a ergonomia moderna do SpellRise.",
        "Nao transformar Prowess em moeda unica sem sinks, caps e protecao de economia.",
        "Nao permitir troca de role/escola em combate se isso comprometer leitura PvP ou balance de TTK.",
        "Nao replicar inventario completo ou payload cosmetico pelo PlayerController.",
        "Nao usar safe zone como decisao client-side; area protegida precisa ser regra server-authoritative.",
    ]:
        bullet(doc, item)

    heading(doc, "13. Fontes consultadas", 1)
    paragraph(doc, "Foram coletadas e resumidas 31 paginas do manual publico. Lista de URLs usadas:")
    for page in pages:
        bullet(doc, f"{page['title']} - {page['url']}")

    doc.save(DOCX_PATH)
    return DOCX_PATH


if __name__ == "__main__":
    path = build_doc()
    print(path)
